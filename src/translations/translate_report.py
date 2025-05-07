import time
from docx import Document
from deep_translator import GoogleTranslator
from concurrent.futures import ThreadPoolExecutor
import os

# Step 1: Load DOCX file
def load_docx_text(filepath):
    doc = Document(filepath)
    # Collecting text from paragraphs
    text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip() != ""])
    
    # Collecting text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Add the text from each cell in the table to the `text`
                text += '\n' + cell.text

    return text
# Step 2: Translate a single chunk of text (thread-safe)
def translate_chunk(chunk, target_lang):
    # Perform the translation and return the result
    translated_chunk = GoogleTranslator(source='auto', target=target_lang).translate(chunk)
    return translated_chunk

# Step 3: Translate text with chunking and ThreadPoolExecutor
def translate_text_in_chunks(text, target_lang, chunk_size=4500, num_threads=None):
    # Split text into chunks of a maximum length of `chunk_size`
    chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

    # Create a ThreadPoolExecutor to handle threads
    if num_threads is None:
        num_threads = os.cpu_count()  # Default to number of CPU cores

    # Use ThreadPoolExecutor to translate chunks in parallel
    with ThreadPoolExecutor(max_workers=num_threads) as executor:
        # Submit translation tasks for each chunk
        futures = [executor.submit(translate_chunk, chunk, target_lang) for chunk in chunks]
        
        # Wait for all the futures to complete and gather the results
        translated_chunks = [future.result() for future in futures]

    # Combine the translated chunks back into a single string
    return '\n'.join(translated_chunks)

# Step 4: Convert to styled HTML
def text_to_html(title, body_text):
    html_template = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <title>{title}</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                line-height: 1.6;
                max-width: 800px;
                margin: 40px auto;
                padding: 20px;
                background-color: #f9f9f9;
                border: 1px solid #ddd;
            }}
            h1 {{
                text-align: center;
                color: #333;
            }}
            p {{
                margin-bottom: 15px;
            }}
        </style>
    </head>
    <body>
        <h1>{title}</h1>
        {body_text}
    </body>
    </html>
    """
    return html_template

# Step 5: Save HTML file
def save_html(content, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(content)

# Step 6: Main function to translate the docx file and measure time
def translate_docx_to_html(docx_path):
    start_time = time.time()  # Start time measurement

    original_text = load_docx_text(docx_path)
    
    translations = {
        'spanish': 'es'
    }

    for lang_name, lang_code in translations.items():
        print(f"Translating to {lang_name.capitalize()}...")
        
        # Perform the translation with chunking and ThreadPoolExecutor
        translated = translate_text_in_chunks(original_text, lang_code)

        # Process the translated text into HTML-friendly paragraphs
        translated_paragraphs = ''.join([f"<p>{para}</p>" for para in translated.split('\n')])

        # Generate HTML content
        html_content = text_to_html(f"Translated Report - {lang_name.capitalize()}", translated_paragraphs)
        
        # Output file name
        output_filename = f"translated_report_{lang_name}.html"
        
        # Save the HTML file
        save_html(html_content, output_filename)
        
        print(f"Saved: {output_filename}")

    # Measure the elapsed time
    elapsed_time = time.time() - start_time
    print(f"Total translation time: {elapsed_time:.2f} seconds")

# Example usage
translate_docx_to_html("../../data/full_reports/Powell ADHD Report - Case 18 - 2021.docx")
