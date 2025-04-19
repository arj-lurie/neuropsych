import PyPDF2
import re
import pdb
import jinja2
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from dateutil import parser

def extract_text_between_pages(text, start_page, end_page):
    # Define the regular expression pattern to match text between 'Page 3' and 'Page 4'
    pattern = rf"{re.escape(start_page)}(.*?){re.escape(end_page)}"
    # pdb.set_trace()
    # Search for the pattern in the text
    match = re.search(pattern, text, re.DOTALL)  # re.DOTALL allows '.' to match newlines
    
    if match:
        # Extract and return the portion of text between 'Page 3' and 'Page 4', without including 'Page 4'
        return match.group(1).strip()
    else:
        return "No match found between the specified pages"

def extract_test_date_and_relationship(text):
    # Use regex to extract Test Date
    test_date_match = re.search(r'Test Date:\s*(\d{1,2}/\d{1,2}/\d{4})', text)
    test_date = test_date_match.group(1) if test_date_match else "Unknown"
    # pdb.set_trace()
    # Determine which label to search for based on text content
    if "Parent Rating Scales" in text:
        relationship_label = "Relationship"
    elif "Teacher Rating Scales" in text:
        relationship_label = "Rater Position"
    else:
        relationship_label = "Relationship"  # fallback/default

    # Use regex to extract Relationship
    relationship_pattern = rf'{relationship_label}:\s*(.*?)\s*(?=Age:|Gender:|Birth Date:|Language:|School:|$)'
    relationship_match = re.search(relationship_pattern, text, re.DOTALL)
    relationship = relationship_match.group(1).strip() if relationship_match else "Unknown"
    if relationship == '':
        relationship = "Unknown"
    
    return {
        'Test Date': test_date,
        'Relationship': relationship,
        'Type': relationship_label
    }

def extract_numbers_and_map_labels(text, labels):

    # Define the patterns
    pattern_general = r"General Combined\s+([\d\s]+)"
    pattern_clinical = r"Clinical Combined\s+([\d\s]+)"

    # Search the text for "General Combined"
    match = re.search(pattern_general, text)

    # If not found, search for "Clinical Combined"
    if not match:
        match = re.search(pattern_clinical, text)
    
    if match:
        # Extract the numbers after "General Combined"
        numbers_str = match.group(1)  # Get the part after "General Combined"
        
        # Split the numbers into a list (assumes numbers are separated by spaces or other whitespace)
        numbers = numbers_str.split()
        
        # Map the numbers to labels (assuming you have enough labels for each number)
        if len(numbers) <= len(labels):
            result = {label: number for label, number in zip(labels, numbers)}
        else:
            result = {label: number for label, number in zip(labels, numbers[:len(labels)])}
        
        return result
    else:
        return "Pattern 'General Combined / Clinical Combined' not found in the text"

def write_to_docx(data, structured_data, fname):
    try:
        # Prepare the values dictionary with the T-scores
        values = {k: v for k, v in data.items()}

        # Create a Word document using python-docx
        doc = Document()

        # Set global line spacing for the document
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing = Pt(10)  # You can adjust the Pt value for your needs

        doc.add_heading('Behavior Assessment System for Children, Third Edition (BASC-3) – Parent Report Form', level=1)
        doc.add_paragraph("Higher scores represent greater endorsed concerns (scores above 60 are considered “at risk” and above 70 are considered clinically significant)")
        
        # Add the table to the Word document
        table = doc.add_table(rows=1, cols=2)

        # Adding the header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Scales'
        hdr_cells[1].text = 'T'

        # Make header cells bold
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = 1  # Center the header text
            # Ensure the 'tcPr' (table cell properties) is initialized
            tcPr = cell._element.get_or_add_tcPr()
            # Add vertical alignment for header row
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

            # Add borders to each header cell
            borders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')  # Border style
                border.set(qn('w:sz'), '4')  # Border size (4 is medium)
                border.set(qn('w:space'), '0')  # Space between borders
                border.set(qn('w:color'), '000000')  # Border color (black)
                borders.append(border)
            tcPr.append(borders)
        
        # Manually add borders to the table (if they don't exist)
        tbl = table._element
        namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        # Find the tblBorders element, or create it if it doesn't exist
        tblBorders = tbl.find('.//w:tblBorders', namespaces=namespace)
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tbl.append(tblBorders)

        # Define border settings
        border_styles = {
            "top": {"val": "single", "sz": "4", "space": "0", "color": "000000"},
            "left": {"val": "single", "sz": "4", "space": "0", "color": "000000"},
            "bottom": {"val": "single", "sz": "4", "space": "0", "color": "000000"},
            "right": {"val": "single", "sz": "4", "space": "0", "color": "000000"},
            "insideH": {"val": "single", "sz": "4", "space": "0", "color": "000000"},
            "insideV": {"val": "single", "sz": "4", "space": "0", "color": "000000"},
        }

        # Apply each border style to the table
        for border, attributes in border_styles.items():
            border_element = tblBorders.find(f'.//w:{border}', namespaces=namespace)
            if border_element is None:
                border_element = OxmlElement(f'w:{border}')
                tblBorders.append(border_element)
            
            for attr, value in attributes.items():
                border_element.set(qn(f'w:{attr}'), value)

        # Add the content rows
        for scale, subscales in structured_data.items():
            # Add the scale row
            row = table.add_row().cells
            row[0].text = scale
            row[1].text = values[scale] + ('*' if scale == 'Adaptive Skills' else '')  # Add '*' for Adaptive Skills

            # Center the "T" and values in the second column
            row[1].paragraphs[0].alignment = 1  # Center the text in the second column

            # Add the subscale rows
            for subscale in subscales:
                # Add '*' to all subscales under Adaptive Skills
                row = table.add_row().cells
                row[0].text = "    " + subscale  # Indented for subscales
                row[1].text = values[subscale] + ('*' if scale == 'Adaptive Skills' or scale == 'Adaptive Skills' else '')

                # Center the "T" and values in the second column
                row[1].paragraphs[0].alignment = 1  # Center the text in the second column
        
        table.style.font.size = Pt(8)  # Optional: You can adjust font size
        table.autofit = False  # Prevent auto resizing
        table.allow_autofit = False  # Disable autofit to maintain consistent spacing

        doc.add_paragraph("*Reverse scored, such that lower scores reflect greater endorsed concerns (scores below 40 are considered “at risk,” and below 30 are considered of significant concern)")
        # Save the document to a file
        doc.save(fname)
        print(f"Document created successfully: {fname}")        
    except Exception as e:
        print(f"Error creating document: {e}: {fname}")    

def read_basc_report(file_path):
    # Read the PDF file and extract text
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
         # Get the number of pages in the PDF
        num_pages = len(reader.pages)
        
        # Extract text from each page
        text = ""
        for page_num in range(num_pages):
            page = reader.pages[page_num]
            text += page.extract_text()

        return text

def create_labels(type):
    if(type == 'Relationship'):
        # Organize the data into hierarchical groups
        structured_data = {
            'Externalizing Problems': ['Hyperactivity', 'Aggression', 'Conduct Problems'],
            'Internalizing Problems': ['Anxiety', 'Depression', 'Somatization'],
            'Behavioral Symptoms Index': ['Attention Problems', 'Atypicality', 'Withdrawal'],
            'Adaptive Skills': ['Adaptability', 'Social Skills', 'Leadership', 'Functional Communication', 'Activities of Daily Living']
        }

        # Create a list of labels for the scales
        labels = [
            "Hyperactivity",
            "Aggression",
            "Conduct Problems",
            "Externalizing Problems",
            "Anxiety",
            "Depression",
            "Somatization",
            "Internalizing Problems",
            "Attention Problems",
            "Atypicality",
            "Withdrawal",
            "Behavioral Symptoms Index",
            "Adaptability",
            "Social Skills",
            "Leadership",
            "Functional Communication",
            "Activities of Daily Living",
            "Adaptive Skills"
        ]
    else:
        # Organize the data into hierarchical groups
        structured_data = {
            'Externalizing Problems': ['Hyperactivity', 'Aggression', 'Conduct Problems'],
            'Internalizing Problems': ['Anxiety', 'Depression', 'Somatization'],
            'School Problems': ['Learning Problems'], #Make sure grouping is not repeated since attention problems is already in the previous set, 
            # but keep it in the labels so mapping during extraction occurs correctly. Also, the reason this is a problem is because the subscale cannot 
            # be repeated across multiple parent scales - unlike adaptability where it does not matter since parent scale is the same.
            'Behavioral Symptoms Index': ['Atypicality', 'Withdrawal'],
            'Adaptive Skills': ['Adaptability', 'Social Skills', 'Leadership', 'Study Skills', 'Functional Communication']
        }

        # Create a list of labels for the scales
        labels = [
            "Hyperactivity",
            "Aggression",
            "Conduct Problems",
            "Externalizing Problems",
            "Anxiety",
            "Depression",
            "Somatization",
            "Internalizing Problems",
            "Attention Problems",
            "Learning Problems",
            "School Problems",
            "Atypicality",
            "Withdrawal",
            "Behavioral Symptoms Index",
            "Adaptability",
            "Social Skills",
            "Leadership",
            "Study Skills",
            "Functional Communication",            
            "Adaptive Skills"
        ]
    return labels, structured_data

def generate_html_table(data, structured_data, rater_label=None):
    try:
        # Prepare the values dictionary with the T-scores
        values = {k: v for k, v in data.items()}
        
        # Function to get background color based on value
        def get_cell_color(value, group):
            try:
                value = int(value)  # Make sure the value is an integer
            except ValueError:
                return ''  # No color for non-numeric values

            # Color coding for Adaptive Skills (Scores < 40)
            if group == "Adaptive" and value <= 30:
                return 'background-color: darkblue; color: white;'  # Clinically significant
            elif group == "Adaptive" and value <= 40:
                return 'background-color: lightblue; color: black;'  # At risk
            elif group != "Adaptive" and value >= 70:
                return 'background-color: red; color: white;'  # Clinically significant
            elif group != "Adaptive" and value >= 60:
                return 'background-color: orange; color: white;'  # At risk

            return ''  # No color for other values

        # Start building HTML
        html = f"""
        <h1>Behavior Assessment System for Children, Third Edition (BASC-3) – Parent Report Form</h1>
        <p>Higher scores represent greater endorsed concerns (scores above 60 are considered “at risk” and above 70 are considered clinically significant)</p>
        <table border="1" cellpadding="6" cellspacing="0" style="border-collapse: collapse; font-family: Arial; font-size: 12px;">
            <thead>
                <tr style="background-color: #f2f2f2;">
                    <th style="text-align: left;">Scales</th>
                    <th style="text-align: center;">{rater_label}</th>
                </tr>
            </thead>
            <tbody>
        """

        # Loop through the structured data for each scale and subscale
        for scale, subscales in structured_data.items():
            # Check if it's an Externalizing/Internalizing scale or Adaptive scale group
            group_type = "External/Internal" if scale in ["Externalizing Problems", "Hyperactivity", "Aggression", "Conduct Problems", 
                                                         "Internalizing Problems", "Anxiety", "Depression", "Somatization", 
                                                         "Behavioral Symptoms Index", "Attention Problems", "Atypicality", 
                                                         "Withdrawal"] else "Adaptive"
            
            # Add the scale row
            t_value = values[scale] + ('*' if scale == 'Adaptive Skills' else '')
            color = get_cell_color(values[scale], group_type)
            html += f"""
                <tr>
                    <td><strong>{scale}</strong></td>
                    <td style="text-align: center; {color}">{t_value}</td>
                </tr>
            """

            # Add subscales
            for subscale in subscales:
                t_sub_value = values[subscale] + ('*' if scale == 'Adaptive Skills' else '')
                color = get_cell_color(values[subscale], group_type)
                html += f"""
                    <tr>
                        <td style="padding-left: 20px;">{subscale}</td>
                        <td style="text-align: center; {color}">{t_sub_value}</td>
                    </tr>
                """

        html += """
            </tbody>
        </table>
        <p>*Reverse scored, such that lower scores reflect greater endorsed concerns (scores below 40 are considered “at risk,” and below 30 are considered of significant concern)</p>
        """

        return html

    except Exception as e:
        print(f"Error generating HTML table: {e}")
        return ""
    
def generate_combined_html_table(all_data, structured_data):
    try:
        def get_cell_color(value, group):
            try:
                value = int(value)
            except (ValueError, TypeError):  # Handle non-numeric or invalid values
                return ''  # Return empty if value is non-numeric

            # Apply color styling based on the group and value thresholds
            if group == "Adaptive" and value <= 30:
                return 'background-color: darkblue; color: white;'
            elif group == "Adaptive" and value <= 40:
                return 'background-color: lightblue; color: black;'
            elif group != "Adaptive" and value >= 70:
                return 'background-color: red; color: white;'
            elif group != "Adaptive" and value >= 60:
                return 'background-color: orange; color: white;'
            return ''  # Default style if no conditions are met

        # Get short name for title (based on the first file)
        first_pdf_name = next(iter(all_data.keys()), 'Unknown PDF')
        short_name = "-".join(first_pdf_name.split("-")[:2]) 

        # Extract rater labels and their associated data
        rater_labels = [(data.get('rater_label', 'Unknown'), pdf_name) for pdf_name, data in all_data.items()]

        # Helper function to extract and standardize date part from the rater_label
        def extract_date(rater_label):
            date_str = rater_label.split(',')[0].strip()
            try:
                parsed_date = parser.parse(date_str)
                return parsed_date.strftime('%Y-%m-%d')
            except ValueError:
                return '1900-01-01'

        # Sort by date descending
        rater_labels_sorted = sorted(rater_labels, key=lambda x: extract_date(x[0]), reverse=True)

        # Create ordered dict of sorted data
        sorted_all_data = {pdf_name: all_data[pdf_name] for _, pdf_name in rater_labels_sorted}
        sorted_rater_labels = [(all_data[pdf_name]['rater_label'], pdf_name) for _, pdf_name in rater_labels_sorted]


        html = f"""
        <h1>{short_name}</h1>
        <p>Higher scores represent greater endorsed concerns (scores above 60 are considered “at risk” and above 70 are considered clinically significant)</p>
        <table border="1" cellpadding="6" cellspacing="0" style="border-collapse: collapse; font-family: Arial; font-size: 12px;">
            <thead>
                <tr style="background-color: #f2f2f2;">
                    <th>Scales/Subscales</th>
        """

        # Add headers
        for rater_label, _ in sorted_rater_labels:
            html += f"<th>{rater_label}</th>"
        html += "</tr></thead><tbody>"

        # Loop over structured_data
        for scale, subscales in structured_data.items():
            group_type = "Adaptive" if scale == "Adaptive Skills" else "External/Internal"

            # Row for main scale
            html += f"<tr><td><strong>{scale}</strong></td>"
            for _, pdf_name in sorted_rater_labels:
                data = sorted_all_data.get(pdf_name)
                val = data['data'].get(scale, "n/a")
                val_display = val + ('*' if group_type == "Adaptive" else '') if val != "n/a" else val
                style = get_cell_color(val, group_type) if val != "n/a" else ''
                html += f"<td style='{style}; text-align:center;'>{val_display}</td>"
            html += "</tr>"

            # Rows for subscales
            for subscale in subscales:
                html += f"<tr><td style='padding-left: 20px;'>{subscale}</td>"
                for _, pdf_name in sorted_rater_labels:
                    data = sorted_all_data.get(pdf_name)
                    val = data['data'].get(subscale, "n/a")
                    val_display = val + ('*' if group_type == "Adaptive" else '') if val != "n/a" else val
                    style = get_cell_color(val, group_type) if val != "n/a" else ''
                    html += f"<td style='{style}; text-align:center;'>{val_display}</td>"
                html += "</tr>"

        html += "</tbody></table><p>*Reverse scored, such that lower scores reflect greater endorsed concerns (scores below 40 are considered “at risk,” and below 30 are considered of significant concern)</p>"
        return html

    except Exception as e:
        print(f"Error generating combined HTML: {e}")
        return ""

